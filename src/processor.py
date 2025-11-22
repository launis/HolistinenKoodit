import os
import time
import google.generativeai as genai
import PyPDF2
import io
import docx

class AIProcessor:
    def __init__(self, api_key=None):
        self.api_key = api_key
        if api_key:
            genai.configure(api_key=api_key)

    def get_available_models(self):
        """Hakee saatavilla olevat mallit."""
        if not self.api_key:
            return ["gemini-flash-latest", "gemini-1.5-flash", "gemini-pro"] # Fallback
        
        try:
            models = []
            for m in genai.list_models():
                if 'generateContent' in m.supported_generation_methods:
                    models.append(m.name.replace("models/", ""))
            return models if models else ["gemini-flash-latest", "gemini-1.5-flash", "gemini-pro"]
        except Exception as e:
            return [f"Virhe: {str(e)}", "gemini-pro"]

    def _read_pdf(self, file_obj):
        """Lukee tekstin PDF-tiedostosta."""
        try:
            pdf_reader = PyPDF2.PdfReader(file_obj)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Virhe luettaessa tiedostoa: {str(e)}"

    def _read_docx(self, file_path_or_obj):
        """Lukee tekstin DOCX-tiedostosta, mukaan lukien taulukot."""
        try:
            doc = docx.Document(file_path_or_obj)
            text = []
            
            # Yksinkertainen tapa: lue kaikki kappaleet
            # Mutta jos dokumentissa on taulukoita, ne jäävät väliin pelkällä doc.paragraphs:lla
            # Joten luetaan "block" kerrallaan jos mahdollista, tai iteroidaan molemmat.
            # python-docx ei tue täydellistä järjestystä helposti, mutta luetaan taulukot lopuksi tai yritetään lomittaa.
            # Yksinkertaisin parannus: Luetaan kappaleet JA taulukot.
            
            for para in doc.paragraphs:
                text.append(para.text)
            
            # Lue myös taulukot
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    # Yhdistä rivin solut | -merkillä jotta rakenne säilyy edes vähän
                    text.append(" | ".join(row_text))
            
            return "\n".join(text)
        except Exception as e:
            return f"Virhe luettaessa DOCX-tiedostoa: {str(e)}"

    def _read_file_content(self, uploaded_file):
        """Lukee tiedoston sisällön tekstinä."""
        if uploaded_file.type == "application/pdf":
            return self._read_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
             return self._read_docx(uploaded_file)
        else:
            # Oletetaan tekstitiedostoksi
            stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
            return stringio.read()

    def process_agent(self, agent_config, uploaded_files, model_name="gemini-pro"):
        """
        Suorittaa tekoälyajon yhdelle agentille.
        """
        if not self.api_key:
            return "Virhe: API-avain puuttuu. Syötä avain asetuksiin."

        agent_name = agent_config.get("name", "Unknown Agent")
        prompt_template = agent_config.get("prompt_template", "")
        
        # 1. Lue tiedostojen sisältö
        files_content = ""
        for uploaded_file in uploaded_files:
            content = self._read_file_content(uploaded_file)
            files_content += f"\n--- TIEDOSTO: {uploaded_file.name} ---\n{content}\n"

        # 2. Rakenna lopullinen kehote
        final_prompt = prompt_template.replace("{files}", files_content)
        
        # 3. Kutsu Gemini API:a
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(final_prompt)
            return response.text
        except Exception as e:
            return f"Virhe tekoälyajossa: {str(e)}"

    def run_batch(self, agents, uploaded_files, model_name="gemini-pro"):
        results = {}
        for agent in agents:
            results[agent["name"]] = self.process_agent(agent, uploaded_files, model_name)
        return results

    def run_orchestration(self, steps, uploaded_files, model_a="gemini-1.5-pro", model_b="gemini-1.5-flash", system_instructions=""):
        """
        Suorittaa monivaiheisen orkestroinnin.
        """
        results = {}
        
        # 1. Lue ladatut tiedostot kerran
        files_content = ""
        for uploaded_file in uploaded_files:
            content = self._read_file_content(uploaded_file)
            files_content += f"\n--- TIEDOSTO: {uploaded_file.name} ---\n{content}\n"

        current_context = f"{system_instructions}\n\n{files_content}"
        
        for step in steps:
            step_id = step["id"]
            task_prompt = step["task_prompt"]
            model_type = step["model_type"] # 'model_a' tai 'model_b'
            
            selected_model = model_a if model_type == "model_a" else model_b
            
            # Lisää edellisten vaiheiden tulokset kontekstiin
            context_with_history = current_context
            if results:
                context_with_history += "\n\n--- AIEMMAT TULOKSET ---\n"
                for key, val in results.items():
                    context_with_history += f"\n=== VAIHE: {key} ===\n{val}\n"
            
            final_prompt = f"{context_with_history}\n\n--- TEHTÄVÄ ---\n{task_prompt}"
            
            try:
                model = genai.GenerativeModel(selected_model)
                response = model.generate_content(final_prompt)
                results[step_id] = response.text
            except Exception as e:
                results[step_id] = f"VIRHE: {str(e)}"
                
        return results
