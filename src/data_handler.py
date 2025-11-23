import io
import docx
import PyPDF2
import re

class TextUpload:
    """Apuluokka tekstin käsittelyyn kuin se olisi ladattu tiedosto."""
    def __init__(self, text, name="pasted_text.txt"):
        self.text = text
        self.name = name
        self.type = "text/plain"
        self._value = text.encode('utf-8')

    def getvalue(self):
        return self._value

class DataHandler:
    """
    Datakerros: Hoitaa tiedostojen lukemisen ja käsittelyn.
    """
    
    def read_file_content(self, uploaded_file):
        """Lukee tiedoston sisällön tekstinä."""
        if uploaded_file.type == "application/pdf":
            return self._read_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
             return self._read_docx(uploaded_file)
        else:
            # Oletetaan tekstitiedostoksi
            stringio = io.StringIO(uploaded_file.getvalue().decode("utf-8"))
            return stringio.read()

    def _read_pdf(self, file_obj):
        """Lukee tekstin PDF-tiedostosta."""
        try:
            reader = PyPDF2.PdfReader(file_obj)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Virhe luettaessa PDF-tiedostoa: {str(e)}"

    def _read_docx(self, file_path_or_obj):
        """Lukee tekstin DOCX-tiedostosta (myös taulukot)."""
        try:
            doc = docx.Document(file_path_or_obj)
            text = []
            for element in doc.element.body:
                if element.tag.endswith('p'): # Paragraph
                    para = docx.text.paragraph.Paragraph(element, doc)
                    text.append(para.text)
                elif element.tag.endswith('tbl'): # Table
                    table = docx.table.Table(element, doc)
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            return f"Virhe luettaessa DOCX-tiedostoa: {str(e)}"

    def parse_prompt_modules(self, file_path_or_obj):
        """
        Parsii Pääarviointikehotteen osiin:
        1. Yleiset säännöt (teksti ennen VAIHE 1:stä)
        2. Vaiheet (VAIHE 1, VAIHE 2, jne.)
        
        Palauttaa: (common_rules, phases_dict)
        """
        try:
            doc = docx.Document(file_path_or_obj)
            common_rules = []
            phases = {}
            current_phase = None
            
            # Regex vaiheen tunnistamiseen: "VAIHE" + välilyöntejä + numero
            phase_pattern = re.compile(r"^(VAIHE\s+\d+)", re.IGNORECASE)
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # Tunnista vaiheen vaihto regexillä
                match = phase_pattern.match(text)
                if match:
                    # Normalisoi avain: "VAIHE 1" (aina isoilla, yksi väli)
                    phase_key = match.group(1).upper().replace("  ", " ")
                    
                    current_phase = phase_key
                    phases[current_phase] = [text] 
                    continue

                if current_phase:
                    phases[current_phase].append(text)
                else:
                    common_rules.append(text)
            
            # Muunna listat tekstiksi
            common_text = "\n".join(common_rules)
            phases_text = {k: "\n".join(v) for k, v in phases.items()}
            
            return common_text, phases_text
            
        except Exception as e:
            return f"Virhe parsiessa kehotetta: {str(e)}", {}
